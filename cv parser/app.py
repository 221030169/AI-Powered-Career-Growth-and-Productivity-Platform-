import streamlit as st
import os
import json
from preprocess_cv import preprocess_cvs
from regex_parser import parse_cv_with_pipeline, CV_FILES_DIR
from resume_ai_analysis import analyze_resume_with_llm   # 👈 NEW import added here

st.set_page_config(page_title="AI Resume Extractor", layout="wide")

st.markdown("""
    <style>
        .big-title {
            font-size: 32px;
            font-weight: bold;
            margin-bottom: 10px;
        }
        .section-title {
            font-size: 22px;
            font-weight: 600;
            margin-top: 20px;
        }
        .content-box {
            background-color: #f0f2f6;
            padding: 15px;
            border-radius: 8px;
            color: #000000;
            margin-bottom: 15px;
            font-family: Arial, sans-serif;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("<div class='big-title'>📄 AI Resume Extractor</div>", unsafe_allow_html=True)
st.write("Upload your resume (PDF, DOC, DOCX) and receive a structured, visually formatted summary of your professional profile.")

uploaded_file = st.file_uploader("Upload your CV", type=['pdf', 'doc', 'docx'])

if uploaded_file:
    saved_path = os.path.join(CV_FILES_DIR, uploaded_file.name)
    with open(saved_path, 'wb') as f:
        f.write(uploaded_file.getbuffer())
    st.success(f"Uploaded: {uploaded_file.name}")

    if st.button("Extract & View Resume"):
        with st.spinner("Processing resume... Please wait ⏳"):
            processed_files = preprocess_cvs()
            all_results = [parse_cv_with_pipeline(file_path) for file_path in processed_files]

        for result in all_results:
            # ---------------- Parsed Resume Display ----------------
            st.markdown(f"<div class='section-title'>👤 {result.get('name', 'Name not found')}</div>", unsafe_allow_html=True)

            st.markdown("<div class='section-title'>📬 Contact Information</div>", unsafe_allow_html=True)
            contact_info = result.get('contact_info', {})
            st.markdown(f"<div class='content-box'>"
                        f"Email: {contact_info.get('email', 'N/A')}<br>"
                        f"Phone: {', '.join(contact_info.get('phone_numbers', [])) or 'N/A'}<br>"
                        f"LinkedIn: {', '.join(contact_info.get('urls', [])) or 'N/A'}"
                        f"</div>", unsafe_allow_html=True)

            st.markdown("<div class='section-title'>🧑‍💻 Skills</div>", unsafe_allow_html=True)
            skills = result.get('skills', [])
            st.markdown(f"<div class='content-box'><ul>{''.join([f'<li>{skill}</li>' for skill in skills])}</ul></div>", unsafe_allow_html=True)

            st.markdown("<div class='section-title'>🎓 Education</div>", unsafe_allow_html=True)
            education_list = result.get('education', [])
            education_html = "<ul>" + "".join([f"<li>{edu}</li>" for edu in education_list]) + "</ul>"
            st.markdown(f"<div class='content-box'>{education_html}</div>", unsafe_allow_html=True)

            st.markdown("<div class='section-title'>🏢 Work Experience</div>", unsafe_allow_html=True)
            experience_list = result.get('experience', [])
            experience_html = "<div class='content-box'>"
            for exp in experience_list:
                if isinstance(exp, dict):
                    role = exp.get('title', '')
                    company = exp.get('company', '')
                    duration = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}".strip(' -')
                    description = exp.get('description', '')
                    experience_html += f"<p><strong>{role}</strong> at {company} <br><em>{duration}</em><br>{description}</p>"
                else:
                    experience_html += f"<p>{exp}</p>"
            experience_html += "</div>"
            st.markdown(experience_html, unsafe_allow_html=True)

            st.markdown("<div class='section-title'>🚀 Projects</div>", unsafe_allow_html=True)
            projects_list = result.get('projects', [])
            projects_html = "<div class='content-box'>"
            for proj in projects_list:
                if isinstance(proj, dict):
                    pname = proj.get('project_name', '')
                    client = proj.get('client_company', '')
                    role = proj.get('role', '')
                    tech = proj.get('technologies_used', '')
                    desc = proj.get('description', '')
                    projects_html += f"<p><strong>{pname}</strong> ({role}) for {client}<br><em>{tech}</em><br>{desc}</p>"
                else:
                    projects_html += f"<p>{proj}</p>"
            projects_html += "</div>"
            st.markdown(projects_html, unsafe_allow_html=True)

            st.markdown("<div class='section-title'>📜 Certifications</div>", unsafe_allow_html=True)
            certs_list = result.get('certifications', [])
            certs_html = "<ul>" + "".join([f"<li>{cert}</li>" for cert in certs_list]) + "</ul>"
            st.markdown(f"<div class='content-box'>{certs_html}</div>", unsafe_allow_html=True)

            st.markdown("<div class='section-title'>🗣️ Languages</div>", unsafe_allow_html=True)
            languages_list = result.get('languages', [])
            languages_html = "<ul>" + "".join([f"<li>{lang}</li>" for lang in languages_list]) + "</ul>"
            st.markdown(f"<div class='content-box'>{languages_html}</div>", unsafe_allow_html=True)

            # ---------------- AI Insights Section (LLM Added) ----------------
            st.markdown("<div class='section-title'>📊 AI Insights (Powered by Ollama)</div>", unsafe_allow_html=True)
            with st.spinner("🤖 Analyzing career growth and ATS score using Ollama..."):
                ai_analysis = analyze_resume_with_llm(result)

            career_score = ai_analysis.get("career_growth_score", "N/A")
            ats_score = ai_analysis.get("ats_score", "N/A")
            recommended_jobs = ai_analysis.get("recommended_jobs", [])
            summary = ai_analysis.get("summary", "")

            st.markdown(f"""
            <div class='content-box'>
            <b>Career Growth Potential:</b> {career_score} / 10<br>
            <b>ATS Compatibility Score:</b> {ats_score} / 100<br>
            <b>Recommended Jobs:</b> {', '.join(recommended_jobs) if recommended_jobs else 'N/A'}<br>
            <b>Summary:</b> {summary}
            </div>
            """, unsafe_allow_html=True)
            # ---------------------------------------------------------------

            # ---------------- Download Buttons ----------------
            json_str = json.dumps(result, indent=4)
            st.download_button(label="📥 Download as JSON", data=json_str,
                               file_name="parsed_resume.json", mime="application/json")

            try:
                import docx
                from io import BytesIO
                doc = docx.Document()
                doc.add_heading('Extracted Resume Summary', 0)

                doc.add_heading('Name', level=1)
                doc.add_paragraph(result.get('name', 'N/A'))

                doc.add_heading('Contact Information', level=1)
                doc.add_paragraph(f"Email: {contact_info.get('email', 'N/A')}\n"
                                  f"Phone: {', '.join(contact_info.get('phone_numbers', []))}\n"
                                  f"LinkedIn: {', '.join(contact_info.get('urls', []))}")

                doc.add_heading('Skills', level=1)
                for skill in skills:
                    doc.add_paragraph(skill, style='List Bullet')

                doc.add_heading('Education', level=1)
                for edu in education_list:
                    doc.add_paragraph(edu, style='List Bullet')

                doc.add_heading('Experience', level=1)
                for exp in experience_list:
                    if isinstance(exp, dict):
                        role = exp.get('title', '')
                        company = exp.get('company', '')
                        duration = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}".strip(' -')
                        description = exp.get('description', '')
                        doc.add_paragraph(f"{role} at {company}\n{duration}\n{description}", style='List Bullet')
                    else:
                        doc.add_paragraph(exp, style='List Bullet')

                doc.add_heading('Projects', level=1)
                for proj in projects_list:
                    if isinstance(proj, dict):
                        pname = proj.get('project_name', '')
                        client = proj.get('client_company', '')
                        role = proj.get('role', '')
                        tech = proj.get('technologies_used', '')
                        desc = proj.get('description', '')
                        doc.add_paragraph(f"{pname} ({role}) for {client}\n{tech}\n{desc}", style='List Bullet')
                    else:
                        doc.add_paragraph(proj, style='List Bullet')

                doc.add_heading('Certifications', level=1)
                for cert in certs_list:
                    doc.add_paragraph(cert, style='List Bullet')

                doc.add_heading('Languages', level=1)
                for lang in languages_list:
                    doc.add_paragraph(lang, style='List Bullet')

                word_file = BytesIO()
                doc.save(word_file)
                word_file.seek(0)

                st.download_button(label="📥 Download as Word (.docx)", data=word_file,
                                   file_name="parsed_resume.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"Word file generation failed: {e}")

            st.success("✅ Resume extraction complete!")
