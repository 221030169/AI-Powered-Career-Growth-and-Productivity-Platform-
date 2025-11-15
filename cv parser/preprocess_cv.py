# Libraries
import os
import re

import fitz
from docx import Document

from config import CV_FILES_DIR, EXTRACTED_TEXT_DIR
from logger import time_function

@time_function
def convert_docx_to_pdf(docx_path):
    """
    Converts a .docx file to .pdf. Requires Microsoft Word and works on Windows.
    Alternatively, can use unoconv for Linux/macOS or a cloud API.
    For simplicity and consistency with existing win32com usage, we'll use a similar approach if available.
    A more robust cross-platform solution would involve LibreOffice/unoconv or third-party libraries.
    """
    try:
        
        from docx2pdf import convert
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        convert(docx_path, pdf_path)
        print(f" [DOCX to PDF] Successfully converted {docx_path} to {pdf_path}")
        return pdf_path
    except ImportError:
        print(" [DOCX to PDF ERROR] 'docx2pdf' library not found. Please install it: pip install docx2pdf")
        print(" [DOCX to PDF INFO] Attempting conversion using win32com.client (Windows only)...")
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
            # 17 = wdFormatPDF
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
            print(f" [DOCX to PDF] Successfully converted {docx_path} to {pdf_path} using win32com.")
            return pdf_path
        except Exception as e:
            print(f" [DOCX to PDF ERROR] Failed to convert {docx_path} to PDF using win32com: {e}")
            return None
    except Exception as e:
        print(f" [DOCX to PDF ERROR] Failed to convert {docx_path} to PDF: {e}")
        return None


def convert_doc_to_docx(doc_path):
    """
    Converts a .doc file to .docx using Word via win32com.client.
    Requires Microsoft Word and works only on Windows.
    """
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        docx_path = doc_path + "x"  
        doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault
        doc.Close()
        word.Quit()
        print(f" [DOC CONVERT] Successfully converted {doc_path} to {docx_path}")
        return docx_path
    except Exception as e:
        print(f" [DOC CONVERT ERROR] Failed to convert {doc_path} to DOCX: {e}")
        return None

# ---- Text Cleaning Function ----
def clean_text(text):
    """
    Performs cleaning on extracted text to remove artifacts and normalize formatting.
    """
    if not isinstance(text, str):
        return ""
    # Normalize line endings
    text = re.sub(r'\r\n|\r','\n',text)
    # Remove common OCR artifacts
    text =re.sub(r'(^|\s)e@(\s|$)', ' ', text)
    text = text.replace('e@', '')
    # Replace specific problematic characters
    text = text.replace('¢', ' ').replace('«', ' ').replace('*', ' ')
    text = text.replace('©', '').replace('™', '').replace('®', '')
    text = text.replace('', '')
    text = text.replace('|', ' ')
    text = text.replace('_', ' ')
    text = text.replace('—', '-')
    text = text.replace('‘', "'").replace('’', "'")
    text = text.replace('“', '"').replace('”', '"')
    text = text.replace('…', '...')
    # Remove bullet points and similar symbols
    text = re.sub(r'[•▪○●✓►‣]', ' ', text)
    text = text.replace('\f','')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text) # Reduce excessive newlines
    # Strip leading/trailing whitespace from each line
    text = '\n'.join([line.strip() for line in text.split('\n')])
    return text.strip()

# ----PDF EXTRACTION FUNC ------
@time_function
def extract_text_from_pdf(pdf_path):
    """
    Extracts both visible text and embedded hyperlinks from a PDF using PyMuPDF (fitz).
    """
    text = ""
    print(f" [PDF DEBUG] Attempting PyMuPDF extraction for PDF: {os.path.basename(pdf_path)}")
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text() + "\n"

            # 🔍 Append any URLs found in hyperlink annotations
            links = page.get_links()
            for link in links:
                uri = link.get("uri", None)
                if uri and uri.startswith("http"):
                    text += f"{uri}\n"  # Append the URL so it's picked up later
        doc.close()

        # if len(text.strip()) < 50:
        #     print(f" [PDF DEBUG] PyMuPDF extraction too little text ({len(text.strip())} chars), falling back to OCR.")
        #     return extract_text_from_pdf_with_ocr(pdf_path)

        print(f" [PDF DEBUG] PyMuPDF extraction successful for {os.path.basename(pdf_path)}")
        return text

    except Exception as e:
        print(f" [PDF DEBUG] PyMuPDF extraction failed for {os.path.basename(pdf_path)}: {type(e).__name__}: {e}. Falling back to OCR.")
        return ""


# ----PDF EXTRACTION WITH OCR FUNC ------
# def extract_text_from_pdf_with_ocr(pdf_path):
#     text = ""
#     print(f" [OCR DEBUG] Attempting OCR for PDF: {os.path.basename(pdf_path)}")
#     try:
#         _poppler_path = globals().get('_poppler_path', None)
#         images = convert_from_path(pdf_path, poppler_path= _poppler_path)
#         for i , image in enumerate(images):
#             print(f"  Performing OCR on page {i+1} of {os.path.basename(pdf_path)}")
#             page_text = pytesseract.image_to_string(image)
#             text += page_text +"\n" # add newline b/w the pages
#         print(f" [OCR DEBUG] OCR extraction complete for: {os.path.basename(pdf_path)}")
#         return text
#     except Exception as e:
#         print(f" [OCR ERROR] OCR extraction failed for {os.path.basename(pdf_path)}: {type(e).__name__}: {e}")
#         return ""
    
# ---- DOCX EXTRACTION FUNC ------
@time_function
def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file, including paragraphs and tables.
    """
    try:
        doc = Document(docx_path)
        full_text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip(): # Only add non-empty paragraphs
                full_text_parts.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            table_text_parts = []
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    # Clean cell text and replace newlines within cells with spaces
                    cell_text = ' '.join(cell.text.strip().splitlines())
                    row_cells_text.append(cell_text)
                # Join cells in a row, e.g., with a tab or ' | ' for clarity
                table_text_parts.append('\t'.join(row_cells_text)) # Using tab for better structure
            # Add table content, separated by newlines, with a clear separator
            if table_text_parts:
                full_text_parts.append("\n--- TABLE START ---\n" + "\n".join(table_text_parts) + "\n--- TABLE END ---")

        return '\n\n'.join(full_text_parts) # Join all parts with double newline
    except Exception as e:
        print(f" [DOCX EXTRACTION ERROR] Failed to extract text from {docx_path}: {e}")
        return None
    
# ---- Main Processing Function ----
@time_function
def preprocess_cvs():
    """
    Scans the CV_FILES, extracts text from DOC, DOCX, and PDF,
    cleans it, applies fallback via .docx → .pdf if necessary,
    and saves the cleaned text to EXTRACTED_TEXT_DIR.
    """
    processed_files_paths = []
    print(f"--- Starting CV preprocessing. Scanning '{CV_FILES_DIR}' ---")

    for filename in os.listdir(CV_FILES_DIR):
        file_path = os.path.join(CV_FILES_DIR, filename)
        if not os.path.isfile(file_path):
            continue

        text = ""
        docx_path = None

        # --- PDF ---
        if filename.endswith(".pdf"):
            print(f" [PDF DETECTED] Extracting text from {filename}...")
            text = extract_text_from_pdf(file_path)

        # --- DOCX ---
        elif filename.endswith(".docx"):
            print(f" [DOCX DETECTED] Extracting text from {filename}...")
            docx_path = file_path
            text = extract_text_from_docx(docx_path)

        # --- DOC ---
        elif filename.endswith(".doc"):
            print(f" [DOC DETECTED] Converting {filename} to .docx...")
            docx_path = convert_doc_to_docx(file_path)
            if docx_path and os.path.exists(docx_path):
                print(f" [DOCX CREATED] Extracting text from converted {docx_path}...")
                text = extract_text_from_docx(docx_path)
            else:
                print(f" [SKIP] Could not convert {filename}. Skipping.")
                continue
        else:
            print(f" [SKIP] Unsupported file type: {filename}.")
            continue

        # --- Fallback to PDF if too short or too long ---
        if docx_path:
            line_count = text.count("\n") + 1
            if line_count <= 6 or len(text) > 131072:
                print(f" ⚠️ Text from {filename} is {'too short' if line_count <= 6 else 'too long'} ({line_count} lines / {len(text)} chars). Trying DOCX→PDF fallback.")
                pdf_path = convert_docx_to_pdf(docx_path)
                if pdf_path and os.path.exists(pdf_path):
                    pdf_text = extract_text_from_pdf(pdf_path)
                    if pdf_text and len(pdf_text.strip()) < len(text.strip()):
                        print(f" ✅ PDF fallback successful. Using extracted text from PDF for {filename}.")
                        text = pdf_text
                    else:
                        print(f" ℹ️ PDF fallback did not improve extraction. Keeping original DOCX text.")
                else:
                    print(f" ❌ Failed to convert {docx_path} to PDF. Keeping DOCX text.")

        # --- Cleaning and Saving ---
        if text and text.strip():
            print(f" [CLEANING] Cleaning text for {filename} ...")
            initial_len = len(text)
            text = clean_text(text)
            cleaned_len = len(text)
            print(f" [CLEANING] Original text length: {initial_len}, Cleaned text length: {cleaned_len}")

            try:
                output_filename = os.path.splitext(filename)[0] + '.txt'
                output_path = os.path.join(EXTRACTED_TEXT_DIR, output_filename)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                processed_files_paths.append(output_path)
                print(f" ✅ Text saved to {output_path}")
            except Exception as e:
                print(f" ❌ Could not save text for {filename}: {type(e).__name__}: {e}")
        else:
            print(f" ⚠️ No meaningful text extracted from {filename}.")

    print(f"\n✅ Finished preprocessing {len(processed_files_paths)} files. Saved in '{EXTRACTED_TEXT_DIR}'.")
    return processed_files_paths

if __name__ == "__main__":
    preprocess_cvs()
